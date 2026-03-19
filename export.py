import os
import tempfile
import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def fix_arabic(text):
    if not text:
        return ""
    reshaped = arabic_reshaper.reshape(str(text))
    return get_display(reshaped)

def export_word(issue, sections):
    doc = Document()

    title = doc.add_heading(issue['newspaper_name'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(f"{issue['title']} — العدد {issue['issue_number']}", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(str(issue['publish_date']))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for section in sections:
        doc.add_heading(section['title'], 2)
        if section.get('image_path') and os.path.exists(section['image_path']):
            try:
                doc.add_picture(section['image_path'], width=Inches(5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
        if section.get('body_text'):
            para = doc.add_paragraph(section['body_text'])
            para.runs[0].font.size = Pt(12)
        doc.add_paragraph()

    footer = doc.add_paragraph(f"Zein Times — @{issue['username']}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    return tmp.name

def export_pdf(issue, sections):
    font_path = os.path.join(os.path.dirname(__file__), "static", "NotoNaskhArabic-Regular.ttf")
    if not os.path.exists(font_path):
        import urllib.request
        os.makedirs(os.path.dirname(font_path), exist_ok=True)
        urllib.request.urlretrieve(
            "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Regular.ttf",
            font_path
        )

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_font("Arabic", "", font_path)

    pdf.set_font("Arabic", size=22)
    pdf.cell(0, 12, fix_arabic(issue['newspaper_name']), ln=True, align="C")
    pdf.set_font("Arabic", size=15)
    pdf.cell(0, 10, fix_arabic(f"{issue['title']} - العدد {issue['issue_number']}"), ln=True, align="C")
    pdf.set_font("Arabic", size=11)
    pdf.cell(0, 8, str(issue['publish_date']), ln=True, align="C")
    pdf.ln(6)

    for section in sections:
        pdf.set_font("Arabic", size=14)
        pdf.cell(0, 10, fix_arabic(section['title']), ln=True)

        if section.get('image_path') and os.path.exists(section['image_path']):
            try:
                pdf.image(section['image_path'], w=180)
                pdf.ln(3)
            except:
                pass

        pdf.set_font("Arabic", size=11)
        if section.get('body_text'):
            pdf.multi_cell(0, 7, fix_arabic(section['body_text']))
        pdf.ln(4)

    pdf.set_font("Arabic", size=10)
    pdf.cell(0, 8, fix_arabic(f"Zein Times — @{issue['username']}"), ln=True, align="C")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    pdf.output(tmp.name)
    return tmp.name